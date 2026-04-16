#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_style(doc, text, level=1):
    """Ajouter un titre avec mise en forme"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_table_with_borders(doc, rows, cols):
    """Ajouter une table avec bordures"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    return table

def shade_cell(cell, color):
    """Mettre en couleur une cellule de tableau"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Créer le document
doc = Document()

# Titre principal
title = doc.add_heading('Home Automation Assistant App', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Sous-titre
subtitle = doc.add_paragraph('Documentation Architecture, Structure et Configuration')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(12)
subtitle_format.font.italic = True

doc.add_paragraph('Version: 1.2.0')
doc.add_paragraph('Date: 25 Mars 2026')
doc.add_paragraph()

# Table des matières
doc.add_heading('Table des Matières', level=1)
toc_items = [
    '1. Vue d\'ensemble du projet',
    '2. Architecture générale',
    '3. Structure des dossiers',
    '4. Configuration (pubspec.yaml)',
    '5. Système de gestion d\'état (Riverpod)',
    '6. Services principaux',
    '7. Écrans et widgets',
    '8. Flux de données MQTT',
    '9. Stockage des préférences',
    '10. Points clés de configuration'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. VUE D'ENSEMBLE
add_heading_with_style(doc, '1. Vue d\'ensemble du projet', level=1)
doc.add_paragraph(
    'Home Automation Assistant App est une application Flutter pour le contrôle et la '
    'surveillance de dispositifs IoT (capteurs et actionneurs) via le protocole MQTT.'
)

table = add_table_with_borders(doc, 3, 2)
table.rows[0].cells[0].text = 'Élément'
table.rows[0].cells[1].text = 'Valeur'
shade_cell(table.rows[0].cells[0], 'E7E6E6')
shade_cell(table.rows[0].cells[1], 'E7E6E6')

table.rows[1].cells[0].text = 'Nom du projet'
table.rows[1].cells[1].text = 'dom_new'
table.rows[2].cells[0].text = 'Version'
table.rows[2].cells[1].text = '1.2.0+1'

doc.add_paragraph()

# 2. ARCHITECTURE GÉNÉRALE
add_heading_with_style(doc, '2. Architecture générale', level=1)

doc.add_heading('Architecture en couches', level=2)
doc.add_paragraph(
    'L\'application suit une architecture en couches avec une séparation claire des responsabilités:'
)

layers = [
    ('Présentation (UI)', 'Écrans et widgets Flutter'),
    ('Gestion d\'état', 'Riverpod providers'),
    ('Services', 'MQTT Service, Preferences Service'),
    ('Données', 'SharedPreferences, MQTT Broker')
]

for layer, desc in layers:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(layer).bold = True
    p.add_run(f': {desc}')

doc.add_paragraph()

doc.add_heading('Diagramme des dépendances', level=2)
doc.add_paragraph('main.dart (MyApp)')
doc.add_paragraph('  ├─ Riverpod ProviderScope', style='List Bullet')
doc.add_paragraph('  ├─ MqttService (Provider global)', style='List Bullet')
doc.add_paragraph('  ├─ PreferencesService (Stockage local)', style='List Bullet')
doc.add_paragraph('  └─ HomeScreen / ConfigScreen', style='List Bullet')

doc.add_paragraph()

# 3. STRUCTURE DES DOSSIERS
add_heading_with_style(doc, '3. Structure des dossiers', level=1)

structure_table = add_table_with_borders(doc, 7, 2)
structure_table.rows[0].cells[0].text = 'Dossier/Fichier'
structure_table.rows[0].cells[1].text = 'Description'
shade_cell(structure_table.rows[0].cells[0], 'E7E6E6')
shade_cell(structure_table.rows[0].cells[1], 'E7E6E6')

folders = [
    ('lib/main.dart', 'Point d\'entrée principal de l\'application'),
    ('lib/providers/app_providers.dart', 'Tous les Riverpod providers'),
    ('lib/services/', 'Services métier (MQTT, Preferences, Startup)'),
    ('lib/screens/', 'Écrans de l\'application'),
    ('lib/widgets/', 'Composants réutilisables'),
    ('android/, ios/, windows/, macos/, linux/', 'Code plateforme spécifique'),
]

for i, (path, desc) in enumerate(folders, 1):
    structure_table.rows[i].cells[0].text = path
    structure_table.rows[i].cells[1].text = desc

doc.add_paragraph()

# 4. CONFIGURATION (pubspec.yaml)
add_heading_with_style(doc, '4. Configuration (pubspec.yaml)', level=1)

doc.add_heading('Informations générales', level=2)
info_table = add_table_with_borders(doc, 4, 2)
info_table.rows[0].cells[0].text = 'Propriété'
info_table.rows[0].cells[1].text = 'Valeur'
shade_cell(info_table.rows[0].cells[0], 'E7E6E6')
shade_cell(info_table.rows[0].cells[1], 'E7E6E6')

info_table.rows[1].cells[0].text = 'Nom'
info_table.rows[1].cells[1].text = 'dom_new'
info_table.rows[2].cells[0].text = 'Version SDK'
info_table.rows[2].cells[1].text = '^3.11.0'
info_table.rows[3].cells[0].text = 'Publiée'
info_table.rows[3].cells[1].text = 'Non (privée)'

doc.add_heading('Dépendances principales', level=2)

deps_table = add_table_with_borders(doc, 9, 3)
deps_table.rows[0].cells[0].text = 'Paquet'
deps_table.rows[0].cells[1].text = 'Version'
deps_table.rows[0].cells[2].text = 'Utilité'
shade_cell(deps_table.rows[0].cells[0], 'E7E6E6')
shade_cell(deps_table.rows[0].cells[1], 'E7E6E6')
shade_cell(deps_table.rows[0].cells[2], 'E7E6E6')

deps = [
    ('flutter_riverpod', '^3.2.1', 'Gestion d\'état'),
    ('riverpod_annotation', '^4.0.2', 'Annotations Riverpod'),
    ('mqtt_client', '^10.11.9', 'Client MQTT'),
    ('device_info_plus', '^12.3.0', 'Infos appareil'),
    ('intl', '^0.20.2', 'Formatage i18n'),
    ('network_info_plus', '^7.0.0', 'Infos réseau'),
    ('collection', '^1.18.0', 'Utilitaires collections'),
    ('shared_preferences', '^2.3.0', 'Stockage local'),
]

for i, (pkg, ver, util) in enumerate(deps, 1):
    deps_table.rows[i].cells[0].text = pkg
    deps_table.rows[i].cells[1].text = ver
    deps_table.rows[i].cells[2].text = util

doc.add_heading('Dev Dependencies', level=2)

dev_table = add_table_with_borders(doc, 5, 3)
dev_table.rows[0].cells[0].text = 'Paquet'
dev_table.rows[0].cells[1].text = 'Version'
dev_table.rows[0].cells[2].text = 'Utilité'
shade_cell(dev_table.rows[0].cells[0], 'E7E6E6')
shade_cell(dev_table.rows[0].cells[1], 'E7E6E6')
shade_cell(dev_table.rows[0].cells[2], 'E7E6E6')

dev_deps = [
    ('build_runner', '^2.4.12', 'Générateur de code'),
    ('riverpod_generator', '^4.0.3', 'Génère code Riverpod'),
    ('json_annotation', '^4.9.0', 'Annotations JSON'),
    ('json_serializable', '^6.8.0', 'Sérialisation JSON'),
]

for i, (pkg, ver, util) in enumerate(dev_deps, 1):
    dev_table.rows[i].cells[0].text = pkg
    dev_table.rows[i].cells[1].text = ver
    dev_table.rows[i].cells[2].text = util

doc.add_page_break()

# 5. RIVERPOD STATE MANAGEMENT
add_heading_with_style(doc, '5. Système de gestion d\'état (Riverpod)', level=1)

doc.add_heading('Principes', level=2)
doc.add_paragraph(
    'Riverpod est utilisé pour gérer l\'état global de l\'application. '
    'Il offre une gestion déclarative et réactive de l\'état.'
)

doc.add_heading('Providers principaux', level=2)

providers_table = add_table_with_borders(doc, 8, 3)
providers_table.rows[0].cells[0].text = 'Provider'
providers_table.rows[0].cells[1].text = 'Type'
providers_table.rows[0].cells[2].text = 'Description'
shade_cell(providers_table.rows[0].cells[0], 'E7E6E6')
shade_cell(providers_table.rows[0].cells[1], 'E7E6E6')
shade_cell(providers_table.rows[0].cells[2], 'E7E6E6')

providers = [
    ('mqttServiceProvider', 'Provider', 'Instance singleton du service MQTT'),
    ('mqttConnectedProvider', 'StateProvider', 'État de connexion MQTT (true/false)'),
    ('temperatureProvider', 'StateProvider.family', 'Température par capteur'),
    ('humiditeProvider', 'StateProvider.family', 'Humidité par capteur'),
    ('actionneurEnabledProvider', 'StateProvider.family', 'Statut activation actionneur'),
    ('actionneurOnOffProvider', 'StateProvider.family', 'État ON/OFF actionneur'),
    ('mqttCapteurMessagesProvider', 'StreamProvider', 'Stream des messages capteurs'),
]

for i, (prov, typ, desc) in enumerate(providers, 1):
    providers_table.rows[i].cells[0].text = prov
    providers_table.rows[i].cells[1].text = typ
    providers_table.rows[i].cells[2].text = desc

doc.add_paragraph()

# 6. SERVICES PRINCIPAUX
add_heading_with_style(doc, '6. Services principaux', level=1)

doc.add_heading('MqttService', level=2)
doc.add_paragraph('Responsabilités:')
doc.add_paragraph('Gestion de la connexion au broker MQTT', style='List Bullet')
doc.add_paragraph('Authentification (username/password)', style='List Bullet')
doc.add_paragraph('Souscription aux topics', style='List Bullet')
doc.add_paragraph('Publication de messages de commande', style='List Bullet')
doc.add_paragraph('Gestion des reconnexions', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Attributs clés:')
mqtt_attrs = [
    'broker: URL du broker MQTT (défaut: 51.77.244.19)',
    'port: Port MQTT (défaut: 1883)',
    'clientId: ID client MQTT unique basé sur le timestamp',
    'username, password: Identifiants d\'authentification',
    'isConnected: État de connexion',
    '_capteurMessagesController: Stream des messages capteurs',
    '_actionneurMessagesController: Stream des commandes actionneurs'
]
for attr in mqtt_attrs:
    doc.add_paragraph(attr, style='List Bullet')

doc.add_paragraph()
doc.add_heading('PreferencesService', level=2)
doc.add_paragraph('Responsabilités:')
doc.add_paragraph('Sauvegarde et chargement de la configuration MQTT', style='List Bullet')
doc.add_paragraph('Utilise SharedPreferences pour la persistence', style='List Bullet')

clés = [
    'mqtt_broker',
    'mqtt_port',
    'mqtt_username',
    'mqtt_password'
]
doc.add_paragraph('Clés de stockage:')
for clé in clés:
    doc.add_paragraph(clé, style='List Bullet')

doc.add_page_break()

# 7. ÉCRANS ET WIDGETS
add_heading_with_style(doc, '7. Écrans et widgets', level=1)

doc.add_heading('Écrans', level=2)
screens_table = add_table_with_borders(doc, 3, 2)
screens_table.rows[0].cells[0].text = 'Écran'
screens_table.rows[0].cells[1].text = 'Fonction'
shade_cell(screens_table.rows[0].cells[0], 'E7E6E6')
shade_cell(screens_table.rows[0].cells[1], 'E7E6E6')

screens_table.rows[1].cells[0].text = 'HomeScreen'
screens_table.rows[1].cells[1].text = 'Affichage principal des capteurs et actionneurs'
screens_table.rows[2].cells[0].text = 'ConfigScreen'
screens_table.rows[2].cells[1].text = 'Configuration des paramètres MQTT'

doc.add_heading('Widgets', level=2)
widgets_table = add_table_with_borders(doc, 6, 2)
widgets_table.rows[0].cells[0].text = 'Widget'
widgets_table.rows[0].cells[1].text = 'Fonction'
shade_cell(widgets_table.rows[0].cells[0], 'E7E6E6')
shade_cell(widgets_table.rows[0].cells[1], 'E7E6E6')

widgets = [
    ('appBar.dart', 'Barre d\'application personnalisée'),
    ('capteur_card.dart', 'Carte affichant les données d\'un capteur'),
    ('actionneur_card.dart', 'Carte de contrôle actionneur simple'),
    ('actionneur_card_2_Btn.dart', 'Carte actionneur avec 2 boutons'),
    ('CouleurCard.dart', 'Carte pour affichage couleur RGB'),
]

for i, (widget, func) in enumerate(widgets, 1):
    widgets_table.rows[i].cells[0].text = widget
    widgets_table.rows[i].cells[1].text = func

doc.add_page_break()

# 8. FLUX DE DONNÉES MQTT
add_heading_with_style(doc, '8. Flux de données MQTT', level=1)

doc.add_heading('Architecture MQTT', level=2)
doc.add_paragraph('L\'application utilise une architecture basée sur les topics MQTT:')

doc.add_heading('Flux entrant (Capteurs)', level=3)
doc.add_paragraph('broker → mqtt_client → StreamController → Riverpod State')

doc.add_heading('Flux sortant (Actionneurs)', level=3)
doc.add_paragraph('UI → MqttService.publish() → broker → appareils')

doc.add_heading('Gestion des messages', level=2)
doc.add_paragraph('Messages reçus:')
doc.add_paragraph('Parsés de JSON', style='List Bullet')
doc.add_paragraph('Envoyés dans le stream correspondant', style='List Bullet')
doc.add_paragraph('Utilisés pour mettre à jour les Providers', style='List Bullet')

doc.add_heading('Lifecycle (Cycle de vie)', level=2)
doc.add_paragraph('1. Au démarrage: LifecycleWrapper écoute le cycle de vie')
doc.add_paragraph('2. onResumed: État de l\'app actif', style='List Bullet')
doc.add_paragraph('3. onPaused: App en arrière-plan', style='List Bullet')
doc.add_paragraph('4. Messages envoyés lors de reprise pour maintenir la synchronisation')

doc.add_page_break()

# 9. STOCKAGE DES PRÉFÉRENCES
add_heading_with_style(doc, '9. Stockage des préférences', level=1)

doc.add_paragraph('L\'application utilise SharedPreferences pour persister la configuration.')

doc.add_heading('Données stockées', level=2)
prefs_table = add_table_with_borders(doc, 5, 3)
prefs_table.rows[0].cells[0].text = 'Clé'
prefs_table.rows[0].cells[1].text = 'Type'
prefs_table.rows[0].cells[2].text = 'Exemple'
shade_cell(prefs_table.rows[0].cells[0], 'E7E6E6')
shade_cell(prefs_table.rows[0].cells[1], 'E7E6E6')
shade_cell(prefs_table.rows[0].cells[2], 'E7E6E6')

prefs = [
    ('mqtt_broker', 'String', '51.77.244.19'),
    ('mqtt_port', 'String', '1883'),
    ('mqtt_username', 'String', 'user123'),
    ('mqtt_password', 'String', 'pass123'),
]

for i, (clé, typ, ex) in enumerate(prefs, 1):
    prefs_table.rows[i].cells[0].text = clé
    prefs_table.rows[i].cells[1].text = typ
    prefs_table.rows[i].cells[2].text = ex

doc.add_paragraph()
doc.add_heading('Emplacement stockage par plateforme', level=2)
platforms = [
    ('Android', '/data/data/com.example.app/shared_prefs/'),
    ('iOS', 'Documents/Library/Preferences/'),
    ('Windows', 'AppData/Roaming/'),
    ('macOS', 'Library/Preferences/'),
    ('Linux', '.config/'),
]

for platform, path in platforms:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(platform).bold = True
    p.add_run(f': {path}')

doc.add_page_break()

# 10. POINTS CLÉS DE CONFIGURATION
add_heading_with_style(doc, '10. Points clés de configuration', level=1)

doc.add_heading('Paramètres d\'environnement', level=2)
doc.add_paragraph('Les paramètres sont externalisés dans PreferencesService:')

doc.add_heading('Configuration initiale', level=2)
config_steps = [
    'L\'app vérifie si config existe au démarrage',
    'Si non: affiche ConfigScreen pour saisir les paramètres',
    'Si oui: charge config et se connecte automatiquement'
]
for i, step in enumerate(config_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Variables d\'environnement', level=2)
settings_table = add_table_with_borders(doc, 5, 3)
settings_table.rows[0].cells[0].text = 'Paramètre'
settings_table.rows[0].cells[1].text = 'Valeur par défaut'
settings_table.rows[0].cells[2].text = 'Configurable'
shade_cell(settings_table.rows[0].cells[0], 'E7E6E6')
shade_cell(settings_table.rows[0].cells[1], 'E7E6E6')
shade_cell(settings_table.rows[0].cells[2], 'E7E6E6')

settings = [
    ('Broker MQTT', '51.77.244.19', 'Oui'),
    ('Port MQTT', '1883', 'Oui'),
    ('Username', 'vide', 'Oui'),
    ('Password', 'vide', 'Oui'),
]

for i, (param, default, config) in enumerate(settings, 1):
    settings_table.rows[i].cells[0].text = param
    settings_table.rows[i].cells[1].text = default
    settings_table.rows[i].cells[2].text = config

doc.add_heading('Reconnexion et résilience', level=2)
doc.add_paragraph('Keep-alive period: 60 secondes', style='List Bullet')
doc.add_paragraph('Auto-reconnexion en cas de déconnexion', style='List Bullet')
doc.add_paragraph('Méthode reconnectWithNewConfig() pour changement de config à l\'runtime', style='List Bullet')

doc.add_page_break()

# Conclusion
add_heading_with_style(doc, 'Conclusion', level=1)
doc.add_paragraph(
    'Home Automation Assistant App 1.2 est une application Flutter bien structurée '
    'utilisant Riverpod pour la gestion d\'état, MQTT pour la communication avec les appareils IoT, '
    'et SharedPreferences pour la persistence. '
    'L\'architecture en couches permet une maintenance facile et une scalabilité.'
)

# Sauvegarder le document
output_path = 'd:\Developpement\Pilotage_Maison_Dev\Home-Assist-App_1.2\ARCHITECTURE_DOCUMENTATION.docx'
doc.save(output_path)
print(f'✓ Document généré: {output_path}')
